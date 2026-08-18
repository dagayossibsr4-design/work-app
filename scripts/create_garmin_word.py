from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

out = "/home/ubuntu/workout-tracker-android/docs/garmin-access-request-template.docx"
doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(10.5)
for style_name in ["Title", "Heading 1", "Heading 2"]:
    styles[style_name].font.name = "Arial"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Garmin Connect Developer Program\n")
r.bold = True
r.font.size = Pt(22)
r2 = p.add_run("תבנית מוכנה לבקשת גישה למפתחים")
r2.bold = True
r2.font.size = Pt(15)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("החלף את הטקסט שבסוגריים המרובעים בפרטים האמיתיים לפני השליחה.").italic = True

def heading(text, level=1):
    doc.add_heading(text, level=level)

def para(text):
    doc.add_paragraph(text)

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(9)

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()

heading("פרטים שכדאי להכין", 1)
table(["פרט", "מה להכין"], [
    ("Applicant / Company", "השם החוקי שלך או שם החברה, אתר ופרטי קשר"),
    ("Product name", "שם האפליקציה באנגלית"),
    ("Product status", "Prototype, beta או production — לפי המצב האמיתי"),
    ("Target platform", "Android / Expo React Native"),
    ("APIs", "Health API ו־Activity API בשלב הראשון"),
    ("Metrics", "Sleep, resting heart rate, stress, Body Battery, steps, calories ופעילות"),
    ("Devices", "Garmin Venu X1 ומכשירי Garmin Connect תואמים"),
    ("Users", "מספר משתמשים צפוי בבטא ובהשקה, בלי לנפח את המספר"),
    ("Privacy", "קישור למדיניות פרטיות, מחיקת נתונים ויצירת קשר"),
    ("Technical contact", "שם, תפקיד, אימייל וטלפון זמינים"),
])

heading("נוסח לתיאור המוצר", 1)
code("""[Product name] is an Android application for real-time workout logging, training-load analysis, and nutrition management. The application helps users record sets, repetitions, weights, exercise history, recovery indicators, and daily nutrition in Hebrew and RTL.

We are requesting Garmin Connect Developer Program access to combine Garmin health and activity data with user-entered workout data. The goal is to provide a more accurate recovery and training-load view, including sleep duration, resting heart rate, stress, Body Battery, and completed activities.

The initial integration will support Garmin Health API and Activity API. We plan to start with a controlled evaluation group and expand only after validating consent, data accuracy, privacy, and synchronization reliability.""")

heading("נוסח למטרת השימוש ב־API", 1)
code("""We intend to use the Health API for user-authorized recovery and all-day health metrics, including sleep, resting heart rate, stress, Body Battery, steps, calories, and respiration where available. We intend to use the Activity API to import completed activity summaries and compare them with the user's manually recorded training sessions.

We will request only the permissions required for these features. Garmin data will not be sold, used for advertising, or shared with unrelated third parties. Users will be able to view the connection status, disconnect Garmin, and request deletion of synchronized Garmin data.""")

heading("נוסח לאבטחה ופרטיות", 1)
code("""OAuth authorization will be completed through Garmin's official authorization page. The Garmin client secret and user access/refresh tokens will never be stored in the Android application or exposed to the client. They will be handled by our server and encrypted at rest using authenticated encryption with key versioning.

The server will associate each Garmin connection with the authenticated application user, validate OAuth state, prevent authorization-code replay, redact secrets from logs, and enforce ownership checks for all synchronized data. We will provide a privacy policy and a user-controlled disconnect and data-deletion flow.""")

heading("נוסח לסביבת בדיקה", 1)
code("""We would like to use the evaluation environment to validate the OAuth flow, consent experience, data synchronization, error handling, token refresh, and data deletion before production release.

Our expected initial test group is approximately [number] users and [number] Garmin devices. The primary device for validation is the Garmin Venu X1, while the application is intended to support other compatible Garmin Connect devices where the requested metrics are available.""")

heading("בקשת APIs", 1)
code("""Requested APIs:
1. Health API — sleep, resting heart rate, stress, Body Battery, steps, calories, respiration, and other approved recovery metrics available for the user's device.
2. Activity API — completed activity summaries and workout-related activity data.

Training API is not required for the initial release because the first version reads Garmin data into the application. We may request Training API separately in a later phase if we decide to publish structured workouts back to Garmin Connect.""")

heading("נוסח סיום", 1)
code("""We would appreciate guidance on the appropriate API permissions, OAuth configuration, evaluation-environment setup, rate limits, and any Garmin-specific branding or privacy requirements. Our technical contact is available for an integration call and can provide additional product, security, or data-flow information upon request.""")

heading("מה לא לכתוב", 1)
para("אין לכתוב שהאפליקציה כבר מחוברת ל־Garmin אם עדיין אין אישור. אין לבקש את כל ה־APIs ללא צורך. אין להציג מספר משתמשים מנופח, ואין להבטיח שמירה או שימוש בנתוני בריאות מעבר למה שהמוצר באמת מבצע. עדיף להציג רשימת metrics מצומצמת ומנומקת.")

heading("קישורים רשמיים", 1)
for link in [
    "Garmin Connect Developer Program: https://developer.garmin.com/gc-developer-program/",
    "Program FAQ: https://developer.garmin.com/gc-developer-program/program-faq/",
    "Health API: https://developer.garmin.com/gc-developer-program/health-api/",
    "Access Request Form: https://www.garmin.com/en-US/forms/GarminConnectDeveloperAccess/",
    "Garmin Developer Contact: https://www.garmin.com/en-US/forms/developercontactus/",
]:
    doc.add_paragraph(link)

doc.save(out)
print(out)
